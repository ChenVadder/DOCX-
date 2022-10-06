from docx import Document
import re

wenjian = Document("D:/lianxi.docx")

for dl in wenjian.paragraphs:
    
    if dl.style.name == 'Heading 1':
        print("Heading 1:",dl.text)
    else:
        print(dl.text)

for dl in wenjian.paragraphs:
    if re.match("^Heading \d+$",dl.style.name):
        print(dl.style.name,": ",dl.text)
    if dl.style.name == "Normal":
        print("正文：",dl.text )
        #此时可以看出，空段落也是一个段落


duan=wenjian.paragraphs[1]

kuai = duan.runs
for k in kuai:
    print(k.text)

wenjian.add_heading("TEST",level=1)
wenjian.save('D:/lianxi2.docx')

wenjian.add_page_break()#分页符
wenjian.add_paragraph("This is Normal Text ")
wenjian.save('D:/lianxi3.docx')

zw = wenjian.paragraphs[-1]
zw.add_run("jiacu").bold =True
zw.add_run("zhengchang")
zw.add_run("xieti").italic = True
wenjian.save('D:/lianxi4.docx')
